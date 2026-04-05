"""
NeuroStride — Pharmacy Router
Ported and integrated from PharmaMind v3 (HackFusion 3)
Full pharmacy: medicines CRUD, orders, prescriptions, 6-agent chat, OCR, analytics, stock logs
"""
import hashlib, io, json, os
from datetime import datetime, timedelta
from typing import Optional, List

import pytesseract
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile
from groq import Groq
from pydantic import BaseModel
from sqlalchemy import Column, Integer, String, Float, Boolean, DateTime, Text, JSON, ForeignKey
from sqlalchemy.orm import Session

import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))
from core.database import get_db, get_current_user, require_role
from models.db_models import User, UserRole, Base

router = APIRouter(prefix="/api/pharmacy2", tags=["Pharmacy Pro"])

GROQ_KEY = os.getenv("GROQ_API_KEY", "").strip()
MODEL    = "llama-3.3-70b-versatile"


# ── Extended DB Models ────────────────────────────────────────────────────────

class Medicine2(Base):
    __tablename__ = "medicines2"
    id           = Column(Integer, primary_key=True, index=True, autoincrement=True)
    name         = Column(String(255), index=True, nullable=False)
    generic_name = Column(String(255), default="")
    brand        = Column(String(255), default="")
    category     = Column(String(100), index=True, default="")
    price        = Column(Float, nullable=False, default=0.0)
    stock        = Column(Integer, default=0)
    threshold    = Column(Integer, default=20)
    rx_required  = Column(Boolean, default=False)
    dosage_form  = Column(String(50), default="tablet")
    strength     = Column(String(50), default="")
    manufacturer = Column(String(255), default="")
    description  = Column(Text, default="")
    side_effects = Column(Text, default="")
    interactions = Column(JSON, default=list)
    expiry       = Column(String(20), default="")
    created_at   = Column(DateTime, default=datetime.utcnow)
    updated_at   = Column(DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)


class Order2(Base):
    __tablename__ = "orders2"
    id             = Column(Integer, primary_key=True, index=True, autoincrement=True)
    order_code     = Column(String(30), unique=True)
    user_id        = Column(String, ForeignKey("users.id"), nullable=True)
    patient_label  = Column(String(255), default="Guest")
    items          = Column(JSON, default=list)
    total          = Column(Float, default=0.0)
    status         = Column(String(50), default="Processing")
    payment_status = Column(String(50), default="Pending")
    razorpay_id    = Column(String(100), nullable=True)
    prescription_id= Column(Integer, nullable=True)
    email_sent     = Column(Boolean, default=False)
    created_at     = Column(DateTime, default=datetime.utcnow)
    updated_at     = Column(DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)


class Prescription2(Base):
    __tablename__ = "prescriptions2"
    id           = Column(Integer, primary_key=True, index=True, autoincrement=True)
    rx_code      = Column(String(30), unique=True)
    patient_id   = Column(String, ForeignKey("users.id"), nullable=True)
    doctor_id    = Column(String, ForeignKey("users.id"), nullable=True)
    patient_name = Column(String(255), default="")
    doctor_name  = Column(String(255), default="")
    items        = Column(JSON, default=list)
    notes        = Column(Text, default="")
    digital_sig  = Column(String(64), default="")
    status       = Column(String(50), default="Active")
    issued_at    = Column(DateTime, default=datetime.utcnow)
    valid_until  = Column(DateTime, default=lambda: datetime.utcnow() + timedelta(days=90))


class StockLog2(Base):
    __tablename__ = "stock_logs2"
    id            = Column(Integer, primary_key=True, autoincrement=True)
    medicine_id   = Column(Integer, ForeignKey("medicines2.id"))
    medicine_name = Column(String(255), default="")
    delta         = Column(Integer, default=0)
    reason        = Column(String(255), default="")
    created_at    = Column(DateTime, default=datetime.utcnow)


# ── Schemas ───────────────────────────────────────────────────────────────────

class MedicineCreate(BaseModel):
    name: str
    generic_name: Optional[str] = ""
    brand: Optional[str] = ""
    category: Optional[str] = ""
    price: float
    stock: int
    threshold: Optional[int] = 20
    rx_required: Optional[bool] = False
    dosage_form: Optional[str] = "tablet"
    strength: Optional[str] = ""
    manufacturer: Optional[str] = ""
    description: Optional[str] = ""
    side_effects: Optional[str] = ""
    interactions: Optional[List[str]] = []
    expiry: Optional[str] = ""

class MedicineUpdate(BaseModel):
    price: Optional[float] = None
    stock: Optional[int] = None
    threshold: Optional[int] = None
    rx_required: Optional[bool] = None
    description: Optional[str] = None
    strength: Optional[str] = None
    category: Optional[str] = None

class OrderCreate(BaseModel):
    items: List[dict]
    patient_label: Optional[str] = "Guest"
    user_id: Optional[str] = None
    prescription_id: Optional[int] = None

class OrderStatusUpdate(BaseModel):
    status: str

class Prescription2Create(BaseModel):
    patient_id: Optional[str] = None
    patient_name: str
    items: List[dict]
    notes: Optional[str] = ""
    validity_days: Optional[int] = 90

class PharmaChatReq(BaseModel):
    message: str
    history: Optional[List[dict]] = []

class StockAdjust(BaseModel):
    delta: int
    reason: Optional[str] = "manual"


# ── Serializers ───────────────────────────────────────────────────────────────

def _med(m: Medicine2) -> dict:
    return {
        "id": m.id, "name": m.name, "generic_name": m.generic_name,
        "brand": m.brand, "category": m.category, "price": m.price,
        "stock": m.stock, "threshold": m.threshold, "rx_required": m.rx_required,
        "dosage_form": m.dosage_form, "strength": m.strength,
        "manufacturer": m.manufacturer, "description": m.description,
        "side_effects": m.side_effects, "interactions": m.interactions or [],
        "expiry": m.expiry,
        "status": "Out of Stock" if m.stock == 0 else "Low Stock" if m.stock <= m.threshold else "In Stock"
    }

def _order(o: Order2) -> dict:
    return {
        "id": o.id, "order_code": o.order_code, "user_id": o.user_id,
        "patient_label": o.patient_label, "items": o.items, "total": o.total,
        "status": o.status, "payment_status": o.payment_status,
        "razorpay_id": o.razorpay_id, "email_sent": o.email_sent,
        "created_at": str(o.created_at)
    }

def _rx(r: Prescription2) -> dict:
    return {
        "id": r.id, "rx_code": r.rx_code, "patient_id": r.patient_id,
        "doctor_id": r.doctor_id, "patient_name": r.patient_name,
        "doctor_name": r.doctor_name, "items": r.items, "notes": r.notes,
        "digital_sig": r.digital_sig, "status": r.status,
        "issued_at": str(r.issued_at), "valid_until": str(r.valid_until)
    }


# ═══════════════════════════════════════════════════════════════════
#  MEDICINE ROUTES
# ═══════════════════════════════════════════════════════════════════

@router.get("/medicines")
def list_medicines(search: str = "", category: str = "", low_stock: bool = False,
                   db: Session = Depends(get_db), _=Depends(get_current_user)):
    q = db.query(Medicine2)
    if search:    q = q.filter(Medicine2.name.ilike(f"%{search}%"))
    if category:  q = q.filter(Medicine2.category == category)
    if low_stock: q = q.filter(Medicine2.stock <= Medicine2.threshold)
    return [_med(m) for m in q.order_by(Medicine2.name).all()]


@router.get("/medicines/search")
def search_medicines(q: str = "", db: Session = Depends(get_db), _=Depends(get_current_user)):
    if len(q) < 2: return []
    return [_med(m) for m in db.query(Medicine2).filter(Medicine2.name.ilike(f"%{q}%")).limit(8).all()]


@router.get("/medicines/categories")
def get_categories(db: Session = Depends(get_db), _=Depends(get_current_user)):
    cats = db.query(Medicine2.category).distinct().all()
    return [c[0] for c in cats if c[0]]


@router.get("/medicines/{mid}")
def get_medicine(mid: int, db: Session = Depends(get_db), _=Depends(get_current_user)):
    m = db.query(Medicine2).filter(Medicine2.id == mid).first()
    if not m: raise HTTPException(404, "Not found")
    return _med(m)


@router.post("/medicines", status_code=201)
def create_medicine(req: MedicineCreate, db: Session = Depends(get_db),
                    _=Depends(require_role(UserRole.PHARMACIST))):
    m = Medicine2(**req.dict())
    db.add(m); db.commit(); db.refresh(m)
    return _med(m)


@router.put("/medicines/{mid}")
def update_medicine(mid: int, req: MedicineUpdate, db: Session = Depends(get_db),
                    _=Depends(require_role(UserRole.PHARMACIST))):
    m = db.query(Medicine2).filter(Medicine2.id == mid).first()
    if not m: raise HTTPException(404)
    for k, v in req.dict(exclude_none=True).items():
        setattr(m, k, v)
    db.commit(); db.refresh(m)
    return _med(m)


@router.delete("/medicines/{mid}")
def delete_medicine(mid: int, db: Session = Depends(get_db),
                    _=Depends(require_role(UserRole.PHARMACIST))):
    m = db.query(Medicine2).filter(Medicine2.id == mid).first()
    if not m: raise HTTPException(404)
    db.delete(m); db.commit()
    return {"deleted": mid}


@router.put("/medicines/{mid}/stock")
def adjust_stock(mid: int, req: StockAdjust, db: Session = Depends(get_db),
                 _=Depends(require_role(UserRole.PHARMACIST))):
    m = db.query(Medicine2).filter(Medicine2.id == mid).first()
    if not m: raise HTTPException(404)
    m.stock = max(0, m.stock + req.delta)
    db.add(StockLog2(medicine_id=mid, medicine_name=m.name, delta=req.delta, reason=req.reason))
    db.commit(); db.refresh(m)
    return _med(m)


# ═══════════════════════════════════════════════════════════════════
#  ORDER ROUTES
# ═══════════════════════════════════════════════════════════════════

@router.post("/orders", status_code=201)
def create_order(req: OrderCreate, db: Session = Depends(get_db),
                 user: Optional[User] = Depends(get_current_user)):
    total = sum(i.get("price", 0) * i.get("qty", 1) for i in req.items)
    code  = f"ORD-{datetime.now().strftime('%y%m%d')}-{db.query(Order2).count()+1:03d}"
    order = Order2(
        order_code=code,
        user_id=user.id if user else req.user_id,
        patient_label=user.full_name if user else req.patient_label,
        items=req.items, total=total, prescription_id=req.prescription_id
    )
    db.add(order); db.flush()
    for item in req.items:
        m = db.query(Medicine2).filter(Medicine2.id == item.get("medicine_id")).first()
        if m:
            qty = item.get("qty", 1)
            m.stock = max(0, m.stock - qty)
            db.add(StockLog2(medicine_id=m.id, medicine_name=m.name, delta=-qty, reason=f"Order {code}"))
    db.commit(); db.refresh(order)
    return _order(order)


@router.get("/orders")
def list_orders(status: str = "", db: Session = Depends(get_db),
                user: Optional[User] = Depends(get_current_user)):
    q = db.query(Order2)
    if user and user.role == UserRole.PATIENT: q = q.filter(Order2.user_id == user.id)
    if status: q = q.filter(Order2.status == status)
    return [_order(o) for o in q.order_by(Order2.created_at.desc()).all()]


@router.put("/orders/{oid}/status")
def update_order_status(oid: int, req: OrderStatusUpdate, db: Session = Depends(get_db),
                        _=Depends(get_current_user)):  # any authenticated user can update
    o = db.query(Order2).filter(Order2.id == oid).first()
    if not o: raise HTTPException(404)
    o.status = req.status
    # Mark payment as Paid when delivered
    if req.status.lower() in ("delivered", "dispensed"):
        o.payment_status = "Paid"
    db.commit()
    return _order(o)


# ═══════════════════════════════════════════════════════════════════
#  PRESCRIPTIONS
# ═══════════════════════════════════════════════════════════════════

@router.post("/prescriptions", status_code=201)
def create_prescription2(req: Prescription2Create, db: Session = Depends(get_db),
                          user: User = Depends(require_role(UserRole.DOCTOR))):
    sig     = hashlib.sha256(f"{req.patient_name}{user.id}{datetime.now().isoformat()}".encode()).hexdigest()
    rx_code = f"RX-{datetime.now().strftime('%y%m%d')}-{db.query(Prescription2).count()+1:03d}"
    rx = Prescription2(
        rx_code=rx_code, patient_id=req.patient_id, doctor_id=user.id,
        patient_name=req.patient_name, doctor_name=user.full_name,
        items=req.items, notes=req.notes, digital_sig=sig[:16], status="Active",
        valid_until=datetime.utcnow() + timedelta(days=req.validity_days or 90)
    )
    db.add(rx); db.commit(); db.refresh(rx)
    return _rx(rx)


@router.get("/prescriptions")
def list_prescriptions2(status: str = "", db: Session = Depends(get_db),
                         user: Optional[User] = Depends(get_current_user)):
    q = db.query(Prescription2)
    if user and user.role == UserRole.PATIENT: q = q.filter(Prescription2.patient_id == user.id)
    if status: q = q.filter(Prescription2.status == status)
    return [_rx(r) for r in q.order_by(Prescription2.issued_at.desc()).all()]


# ═══════════════════════════════════════════════════════════════════
#  BILL GENERATION
# ═══════════════════════════════════════════════════════════════════

@router.get("/orders/{order_id}/bill")
def download_bill2(
    order_id: int, db: Session = Depends(get_db),
    _=Depends(get_current_user)
):
    """Generate Word document bill for a pharmacy2 order."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io
    from fastapi.responses import StreamingResponse

    order = db.query(Order2).filter(Order2.id == order_id).first()
    if not order: raise HTTPException(404, "Order not found")

    doc = Document()
    title = doc.add_heading("NeuroStride PharmaMind", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(252, 109, 38)
    sub = doc.add_paragraph("AI-Powered Pharmacy · Ideathon LPU 2026")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_heading("Tax Invoice", 1)
    meta = doc.add_table(rows=3, cols=2)
    meta.style = "Table Grid"
    for i, (k, v) in enumerate([
        ("Invoice No.", f"INV-{str(order.id).upper()[:8]}"),
        ("Order Code",  order.order_code or str(order.id)),
        ("Date",        str(order.created_at)[:10]),
    ]):
        meta.rows[i].cells[0].text = k
        meta.rows[i].cells[1].text = str(v)
        meta.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    doc.add_heading("Medicines", 1)
    items = order.items or []
    if items:
        t = doc.add_table(rows=1 + len(items), cols=4)
        t.style = "Table Grid"
        for j, h in enumerate(["Medicine", "Qty", "Unit Price (₹)", "Total (₹)"]):
            t.rows[0].cells[j].text = h
            t.rows[0].cells[j].paragraphs[0].runs[0].bold = True
        subtotal = 0
        for i, item in enumerate(items):
            qty = item.get("qty", 1); price = item.get("price", 0); total = qty * price
            subtotal += total
            t.rows[i+1].cells[0].text = item.get("name", "?")
            t.rows[i+1].cells[1].text = str(qty)
            t.rows[i+1].cells[2].text = f"Rs.{price:.2f}"
            t.rows[i+1].cells[3].text = f"Rs.{total:.2f}"
    else:
        subtotal = order.total or 0
        doc.add_paragraph("No items recorded.")

    doc.add_paragraph()
    service_fee = round(subtotal * 0.02, 2)
    gst = round(subtotal * 0.18, 2)
    grand = round(subtotal + service_fee + gst, 2)
    totals = doc.add_table(rows=4, cols=2)
    totals.style = "Table Grid"
    for i, (k, v) in enumerate([
        ("Subtotal", f"Rs.{subtotal:.2f}"),
        ("Service Charge (2%)", f"Rs.{service_fee:.2f}"),
        ("GST @ 18%", f"Rs.{gst:.2f}"),
        ("GRAND TOTAL", f"Rs.{grand:.2f}"),
    ]):
        totals.rows[i].cells[0].text = k
        totals.rows[i].cells[1].text = v
        if i == 3:
            totals.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            totals.rows[i].cells[1].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    footer = doc.add_paragraph("Thank you for choosing NeuroStride PharmaMind")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    return StreamingResponse(buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=Bill_{order.order_code}.docx"})


# ═══════════════════════════════════════════════════════════════════
#  ANALYTICS + STOCK LOGS
# ═══════════════════════════════════════════════════════════════════

@router.get("/analytics")
def analytics(db: Session = Depends(get_db),
              _=Depends(require_role(UserRole.PHARMACIST, UserRole.DOCTOR))):
    orders = db.query(Order2).all()
    meds   = db.query(Medicine2).all()
    total_rev   = sum(o.total or 0 for o in orders if (o.status or "").lower() in ("delivered", "dispensed"))
    low_stock   = [_med(m) for m in meds if m.stock <= m.threshold]
    status_counts: dict = {}
    for o in orders:
        status_counts[o.status] = status_counts.get(o.status, 0) + 1
    name_counts: dict = {}
    for o in orders:
        for item in (o.items or []):
            n = item.get("name", ""); name_counts[n] = name_counts.get(n, 0) + item.get("qty", 1)
    trending = sorted(name_counts.items(), key=lambda x: x[1], reverse=True)[:10]
    top = [{"name": n, "sold": q} for n, q in trending]
    return {
        "total_revenue": total_rev, "total_orders": len(orders),
        "total_medicines": len(meds), "low_stock_alerts": low_stock,
        "status_breakdown": status_counts,
        "trending_medicines": top,
        "top_medicines": top,
        "recent_orders": [_order(o) for o in sorted(orders, key=lambda x: x.created_at, reverse=True)[:5]]
    }


@router.get("/stock-logs")
def stock_logs(db: Session = Depends(get_db), _=Depends(require_role(UserRole.PHARMACIST))):
    logs = db.query(StockLog2).order_by(StockLog2.created_at.desc()).limit(100).all()
    return [{"id": l.id, "medicine_name": l.medicine_name, "delta": l.delta,
             "reason": l.reason, "created_at": str(l.created_at)} for l in logs]


# ═══════════════════════════════════════════════════════════════════
#  6-AGENT PHARMACY CHAT
# ═══════════════════════════════════════════════════════════════════

@router.post("/chat")
async def pharmacy_chat(req: PharmaChatReq, db: Session = Depends(get_db),
                         user: Optional[User] = Depends(get_current_user)):
    meds = db.query(Medicine2).all()

    # Full medicine context — same as original PharmaMind
    meds_ctx = "\n".join([
        f"- ID:{m.id} | {m.name} | {m.generic_name or ''} | {m.category} | "
        f"₹{m.price} | Stock:{m.stock} | Rx:{m.rx_required} | "
        f"Form:{m.dosage_form} | Strength:{m.strength} | "
        f"Interactions:{','.join(m.interactions or [])}"
        for m in meds
    ])

    # User context
    user_allergies  = []
    user_conditions = []
    user_orders_ctx = []
    if user:
        recent_orders = db.query(Order2).filter(
            Order2.user_id == user.id
        ).order_by(Order2.created_at.desc()).limit(5).all()
        user_orders_ctx = [{"id": o.order_code, "items": o.items} for o in recent_orders]

    system = f"""You are PharmaMind — an autonomous AI pharmacist powered by 6 specialist agents.
You MUST respond ONLY with valid JSON. Do NOT add any text before or after the JSON.

=== MEDICINE DATABASE (REAL-TIME) ===
{meds_ctx}

=== PATIENT CONTEXT ===
Allergies: {user_allergies}
Conditions: {user_conditions}
Recent orders: {json.dumps(user_orders_ctx)}

=== 6-AGENT SYSTEM ===
Agent 1 — Order Processing: Extracts medicine name, dosage, quantity from natural language. Understands symptoms and maps them to medicines.
Agent 2 — Safety Validation: Checks drug interactions, allergy conflicts, Rx requirements. NEVER approves if there is an allergy conflict.
Agent 3 — Refill Intelligence: Detects refill needs from order history patterns.
Agent 4 — Inventory Monitor: Checks real-time stock levels, suggests alternatives if out of stock.
Agent 5 — Prescription Verify: Validates Rx requirement before allowing order.
Agent 6 — Analytics & Fraud: Detects unusual ordering patterns.

=== RESPONSE FORMAT (return ONLY this JSON, nothing else) ===
{{
  "agent1": {{"medicine_name":"","dosage":"","quantity":1,"intent":"order|query|symptom|refill|general","detected_symptoms":[]}},
  "agent2": {{"prescription_required":false,"allergy_warning":null,"interaction_warning":null,"approved":true,"reason":""}},
  "agent3": {{"refill_needed":false,"days_until_refill":null,"recommendation":""}},
  "agent4": {{"in_stock":true,"stock_level":0,"alternatives":[],"price":0}},
  "agent5": {{"rx_check_passed":true,"message":""}},
  "agent6": {{"fraud_flag":false,"reason":""}},
  "response": "Friendly conversational reply with emojis explaining what agents found. Mention medicine names, prices, stock status.",
  "suggested_medicines": [{{"id":0,"name":"","qty":1,"price":0,"rx_required":false,"strength":""}}],
  "action": "add_to_cart|show_info|request_prescription|out_of_stock|none"
}}

RULES:
- Respond in the SAME language the user writes in (Hindi/English/etc)
- For symptom queries: suggest appropriate medicines from the database, set action to add_to_cart
- For out of stock: set action to out_of_stock and suggest alternatives
- For Rx medicines: set agent2.prescription_required=true and action=request_prescription
- Always populate suggested_medicines when recommending anything
- Use emojis in the response to make it friendly 💊🛒✅"""

    msgs = [{"role": h["role"], "content": h["content"]}
            for h in (req.history or [])[-6:]
            if h.get("role") in ("user", "assistant")]
    msgs.append({"role": "user", "content": req.message})

    try:
        client = Groq(api_key=GROQ_KEY)
        completion = client.chat.completions.create(
            model=MODEL, max_tokens=1800, temperature=0.3,
            messages=[{"role": "system", "content": system}] + msgs
        )
        raw   = completion.choices[0].message.content.strip()
        start = raw.find("{"); end = raw.rfind("}") + 1
        return json.loads(raw[start:end])
    except json.JSONDecodeError:
        # Try to extract response text even if JSON is malformed
        raw = completion.choices[0].message.content.strip() if 'completion' in dir() else ""
        return {
            "response": raw or "I processed your request.",
            "suggested_medicines": [], "action": "none",
            "agent1": {}, "agent2": {"approved": True},
            "agent3": {}, "agent4": {}, "agent5": {}, "agent6": {}
        }
    except Exception as e:
        return {
            "response": f"AI error: {str(e)}. Please check GROQ_API_KEY.",
            "suggested_medicines": [], "action": "none",
            "agent1": {}, "agent2": {"approved": True},
            "agent3": {}, "agent4": {}, "agent5": {}, "agent6": {}
        }


# ═══════════════════════════════════════════════════════════════════
#  OCR — PharmaMind's enhanced version with --psm 6
# ═══════════════════════════════════════════════════════════════════

@router.post("/ocr")
async def ocr_prescription2(file: UploadFile = File(...), db: Session = Depends(get_db),
                              _=Depends(get_current_user)):
    img_bytes = await file.read()
    if not img_bytes: raise HTTPException(400, "Empty file")

    raw_text = ""
    try:
        from PIL import Image
        img = Image.open(io.BytesIO(img_bytes))
        if img.mode not in ("RGB", "L"): img = img.convert("RGB")
        raw_text = pytesseract.image_to_string(img, lang="eng", config="--psm 6").strip()
    except Exception as e:
        raw_text = f"[OCR error: {e}]"

    if not raw_text or raw_text.startswith("[OCR error"):
        return {"doctor_name": "", "patient_name": "", "date": "", "medicines": [], "raw_ocr": raw_text}

    meds      = db.query(Medicine2).all()
    med_names = [m.name for m in meds]
    safe_raw  = raw_text[:300].replace('"', "'").replace('\n', ' ').replace('\\', '')

    parse_prompt = f"""Parse prescription OCR text. Use medical knowledge to correct OCR errors.
Match medicine names to database when possible. Return ONLY valid JSON:
{{"doctor_name":"","patient_name":"","date":"","medicines":[{{"name":"","dosage":"","quantity":1,"duration":"","in_database":false}}],"raw_ocr":"{safe_raw}","confidence":0.0}}

OCR TEXT: {raw_text}
Known medicines: {med_names[:30]}"""

    try:
        client = Groq(api_key=GROQ_KEY)
        completion = client.chat.completions.create(model=MODEL, temperature=0.1, max_tokens=800,
            messages=[{"role": "system", "content": "Medical prescription parser. Return valid JSON only. No markdown."},
                      {"role": "user", "content": parse_prompt}])
        parsed = completion.choices[0].message.content.strip()
        clean  = parsed.replace("```json","").replace("```","").strip()
        start  = clean.find("{"); end = clean.rfind("}") + 1
        result = json.loads(clean[start:end])
        result["raw_ocr"] = raw_text
        return result
    except Exception as e:
        return {"doctor_name": "", "patient_name": "", "date": "", "medicines": [],
                "raw_ocr": raw_text, "confidence": 0.0, "error": str(e)}
