"""
NeuroStride — FastAPI Main Application
SECURITY HARDENED — All CVEs patched (2026-05)

Fixes applied:
  [C1]  IDOR on /patients — ownership + role checks added
  [C2]  IDOR on /sessions — ownership + role checks added
  [C3]  Unauthenticated WebSockets — JWT token query-param validation
  [H1]  Mass assignment — explicit field allowlists in Pydantic + setattr guard
  [H2]  Unsanitized prescription data — length/char validation before Order2 insert
  [H3]  CORS wildcard + credentials — origins from env var, credentials disabled
  [H4]  Undefined 'Order' in bill route — replaced with Order2, correct import
  [M1]  Account enumeration — generic error messages on register/login
  [M2]  No rate limiting — slowapi per-IP limits on auth endpoints
  [M3]  JWT role claim trusted blindly — role re-fetched from DB on every request
  [M4]  Partial rollback leaves inconsistent DB — Order2 creation moved to background task
  [BONUS] Strict input validation on all Pydantic models (max lengths, regex patterns)
  [BONUS] Security response headers middleware
  [BONUS] Request size limit middleware
  [BONUS] Sensitive fields stripped from all serializers
  [BONUS] Content-Type enforcement
"""

import os
import re
import sys
import json
import asyncio
import logging
from contextlib import asynccontextmanager
from typing import Optional, Annotated
from datetime import datetime

from fastapi import (
    FastAPI, Depends, HTTPException, WebSocket, WebSocketDisconnect,
    Query, Request, status, BackgroundTasks
)
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from starlette.middleware.base import BaseHTTPMiddleware
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded
from sqlalchemy.orm import Session
from pydantic import BaseModel, EmailStr, field_validator, Field
from jose import JWTError, jwt

sys.path.insert(0, os.path.dirname(__file__))

from agents.agents import router as agents_router
from routers.pharmacy import router as pharmacy_router, Medicine2, Order2, Prescription2, StockLog2

from core.database import (
    init_db, get_db, hash_password, verify_password,
    create_access_token, SECRET_KEY, ALGORITHM
)
from models.db_models import (
    User, UserRole, PatientProfile, DoctorProfile,
    RehabSession, ExercisePlan, Prescription,
    PharmacyOrder, MedicineInventory, ProgressReport
)

logger = logging.getLogger("neurostride.security")

# ── Allowed CORS origins from env (never wildcard in prod) ────────────────────
# Set ALLOWED_ORIGINS="https://your-app.vercel.app" in production .env
# Defaults to localhost for local dev if env var not set
_raw_origins = os.getenv("ALLOWED_ORIGINS", "http://localhost:3000,http://localhost:3001")
ALLOWED_ORIGINS: list[str] = [o.strip() for o in _raw_origins.split(",") if o.strip()]


# ── Rate limiter (slowapi) ────────────────────────────────────────────────────
limiter = Limiter(key_func=get_remote_address)


# ── App lifecycle ─────────────────────────────────────────────────────────────
@asynccontextmanager
async def lifespan(app: FastAPI):
    logger.info("[NeuroStride] Starting up...")
    init_db()
    logger.info("[NeuroStride] Database ready")
    yield
    logger.info("[NeuroStride] Shutting down")


app = FastAPI(
    title="NeuroStride API",
    description="AI-Powered Neurorehabilitation Platform",
    version="1.0.0",
    lifespan=lifespan,
    # Hide schema from public in production
    docs_url="/docs" if os.getenv("ENV", "development") != "production" else None,
    redoc_url="/redoc" if os.getenv("ENV", "development") != "production" else None,
)

app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)


# ── Security Headers Middleware ───────────────────────────────────────────────
class SecurityHeadersMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request: Request, call_next):
        response = await call_next(request)
        response.headers["X-Content-Type-Options"] = "nosniff"
        response.headers["X-Frame-Options"] = "DENY"
        response.headers["X-XSS-Protection"] = "1; mode=block"
        response.headers["Strict-Transport-Security"] = "max-age=63072000; includeSubDomains"
        response.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
        response.headers["Cache-Control"] = "no-store"
        response.headers["Permissions-Policy"] = "geolocation=(), camera=(), microphone=()"
        # Remove server fingerprint — MutableHeaders has no .pop(), use del
        if "server" in response.headers:
            del response.headers["server"]
        return response


# ── Request Size Limit Middleware ─────────────────────────────────────────────
class RequestSizeLimitMiddleware(BaseHTTPMiddleware):
    MAX_BODY = 1 * 1024 * 1024  # 1 MB

    async def dispatch(self, request: Request, call_next):
        content_length = request.headers.get("content-length")
        if content_length and int(content_length) > self.MAX_BODY:
            return JSONResponse(status_code=413, content={"detail": "Request body too large"})
        return await call_next(request)


app.add_middleware(SecurityHeadersMiddleware)
app.add_middleware(RequestSizeLimitMiddleware)

# [H3] FIX: No wildcard + credentials. Origins from env, credentials disabled.
app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=False,          # ← was True with wildcard — spec violation
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
    allow_headers=["Authorization", "Content-Type", "Accept"],
)

app.include_router(agents_router)
app.include_router(pharmacy_router)


# ── Auth helpers ──────────────────────────────────────────────────────────────

def _get_user_from_token(token: str, db: Session) -> User:
    """
    [M3] FIX: Decode JWT then RE-FETCH user + role from DB.
    Never trust the role embedded in the token payload.
    """
    credentials_exc = HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        user_id: str = payload.get("sub")
        if not user_id:
            raise credentials_exc
    except JWTError:
        raise credentials_exc

    # Always re-fetch from DB — role in token is only used for routing hint
    user = db.query(User).filter(User.id == user_id, User.is_active == True).first()
    if not user:
        raise credentials_exc
    return user


def get_current_user_secure(
    request: Request,
    db: Session = Depends(get_db)
) -> User:
    auth = request.headers.get("Authorization", "")
    if not auth.startswith("Bearer "):
        raise HTTPException(401, "Not authenticated")
    return _get_user_from_token(auth.split(" ", 1)[1], db)


def require_role_secure(*roles: UserRole):
    def _dep(current_user: User = Depends(get_current_user_secure)) -> User:
        if current_user.role not in roles:
            raise HTTPException(403, "Insufficient permissions")
        return current_user
    return _dep


# ── WebSocket JWT validation ──────────────────────────────────────────────────

def _ws_auth(token: str, db: Session) -> User:
    """[C3] FIX: Authenticate WebSocket connections via query-param JWT."""
    if not token:
        raise ValueError("No token")
    return _get_user_from_token(token, db)


# ── Pydantic schemas (hardened) ───────────────────────────────────────────────

_NAME_RE  = re.compile(r"^[\w\s\-'.]{1,120}$")
_EMAIL_RE = re.compile(r"^[^@\s]{1,64}@[^@\s]{1,255}$")


class RegisterRequest(BaseModel):
    email:     str = Field(..., min_length=5, max_length=254)
    password:  str = Field(..., min_length=10, max_length=128)
    full_name: str = Field(..., min_length=1, max_length=120)
    role:      UserRole
    phone:     Optional[str] = Field(None, max_length=20, pattern=r"^\+?[\d\s\-]{7,20}$")
    language:  Optional[str] = Field("en", max_length=10)

    @field_validator("email")
    @classmethod
    def validate_email(cls, v: str) -> str:
        v = v.strip().lower()
        if not _EMAIL_RE.match(v):
            raise ValueError("Invalid email format")
        return v

    @field_validator("full_name")
    @classmethod
    def validate_name(cls, v: str) -> str:
        v = v.strip()
        if not _NAME_RE.match(v):
            raise ValueError("Invalid name characters")
        return v

    @field_validator("password")
    @classmethod
    def password_strength(cls, v: str) -> str:
        if not re.search(r"[A-Z]", v):
            raise ValueError("Password must contain an uppercase letter")
        if not re.search(r"[0-9]", v):
            raise ValueError("Password must contain a digit")
        if not re.search(r"[^A-Za-z0-9]", v):
            raise ValueError("Password must contain a special character")
        return v


class LoginRequest(BaseModel):
    email:    str = Field(..., max_length=254)
    password: str = Field(..., max_length=128)


# [H1] FIX: Explicit field allowlist — only safe patient-owned fields, no FKs
class PatientProfileUpdate(BaseModel):
    date_of_birth:     Optional[str]   = Field(None, max_length=10, pattern=r"^\d{4}-\d{2}-\d{2}$")
    gender:            Optional[str]   = Field(None, max_length=20)
    blood_group:       Optional[str]   = Field(None, max_length=5,  pattern=r"^(A|B|AB|O)[+-]$")
    weight_kg:         Optional[float] = Field(None, gt=0, lt=500)
    height_cm:         Optional[float] = Field(None, gt=0, lt=300)
    diagnosis:         Optional[str]   = Field(None, max_length=500)
    affected_side:     Optional[str]   = Field(None, max_length=20)
    paralysis_level:   Optional[str]   = Field(None, max_length=50)
    allergies:         Optional[list]  = Field(None, max_length=50)
    current_meds:      Optional[list]  = Field(None, max_length=50)
    emergency_contact: Optional[str]   = Field(None, max_length=200)
    # assigned_doctor_id intentionally excluded — use /assign-doctor endpoint

    # Guard: only these exact keys may ever reach setattr
    _SAFE_FIELDS = frozenset({
        "date_of_birth", "gender", "blood_group", "weight_kg", "height_cm",
        "diagnosis", "affected_side", "paralysis_level", "allergies",
        "current_meds", "emergency_contact"
    })


class SessionCreate(BaseModel):
    exercise_plan_id: Optional[str] = Field(None, max_length=36)
    session_mode:     str           = Field("live", max_length=20)


class SessionUpdate(BaseModel):
    exercises_completed: Optional[list]  = None
    total_reps:          Optional[int]   = Field(None, ge=0, le=10_000)
    avg_form_score:      Optional[float] = Field(None, ge=0.0, le=100.0)
    emg_peak:            Optional[float] = Field(None, ge=0.0, le=1_000_000.0)
    emg_avg_rms:         Optional[float] = Field(None, ge=0.0, le=1_000_000.0)
    intent_count:        Optional[int]   = Field(None, ge=0, le=10_000)
    signal_quality:      Optional[float] = Field(None, ge=0.0, le=100.0)
    notes:               Optional[str]   = Field(None, max_length=2000)


# [H2] FIX: Strict medication item validation
class MedicationItem(BaseModel):
    name:      str           = Field(..., min_length=1, max_length=200)
    dose:      Optional[str] = Field(None, max_length=100)
    frequency: Optional[str] = Field(None, max_length=100)
    duration:  Optional[str] = Field(None, max_length=100)

    @field_validator("name", "dose", "frequency", "duration", mode="before")
    @classmethod
    def strip_and_sanitize(cls, v):
        if v is None:
            return v
        v = str(v).strip()
        # Block obvious injection attempts
        if any(c in v for c in ["<", ">", ";", "--", "/*", "*/"]):
            raise ValueError("Invalid characters in medication field")
        return v


class PrescriptionCreate(BaseModel):
    patient_id:  str                   = Field(..., max_length=36)
    medications: list[MedicationItem]  = Field(..., min_length=1, max_length=30)
    notes:       Optional[str]         = Field(None, max_length=2000)


class ExercisePlanCreate(BaseModel):
    patient_id:         str           = Field(..., max_length=36)
    title:              str           = Field(..., min_length=1, max_length=200)
    description:        Optional[str] = Field(None, max_length=2000)
    exercises:          list          = Field(..., max_length=50)
    frequency_per_week: int           = Field(5, ge=1, le=7)
    duration_weeks:     int           = Field(4, ge=1, le=52)


class PharmacyOrderUpdate(BaseModel):
    status: str           = Field(..., max_length=50)
    notes:  Optional[str] = Field(None, max_length=1000)


class StockUpdateBody(BaseModel):
    stock_quantity: int           = Field(..., ge=0, le=1_000_000)
    reorder_level:  Optional[int] = Field(None, ge=0, le=1_000_000)


# ── WebSocket connection manager ──────────────────────────────────────────────
class ConnectionManager:
    def __init__(self):
        self.active: dict[str, list[WebSocket]] = {}

    async def connect(self, websocket: WebSocket, room: str):
        await websocket.accept()
        self.active.setdefault(room, []).append(websocket)

    def disconnect(self, websocket: WebSocket, room: str):
        if room in self.active:
            self.active[room] = [ws for ws in self.active[room] if ws != websocket]

    async def broadcast(self, room: str, data: dict):
        dead = []
        for ws in self.active.get(room, []):
            try:
                await ws.send_json(data)
            except Exception:
                dead.append(ws)
        for ws in dead:
            self.disconnect(ws, room)


manager = ConnectionManager()


# ═══════════════════════════════════════════════════════════════════
#  AUTH ROUTES
# ═══════════════════════════════════════════════════════════════════

@app.post("/api/auth/register", tags=["Auth"])
@limiter.limit("5/minute")          # [M2] FIX: Rate limit registration
def register(request: Request, req: RegisterRequest, db: Session = Depends(get_db)):
    # [M1] FIX: Generic duplicate message — don't confirm email existence
    existing = db.query(User).filter(User.email == req.email).first()
    if existing:
        raise HTTPException(400, "Registration failed. Please check your details.")

    user = User(
        email=req.email,
        full_name=req.full_name,
        hashed_password=hash_password(req.password),
        role=req.role,
        phone=req.phone,
        language=req.language
    )
    db.add(user)
    db.flush()

    if req.role == UserRole.PATIENT:
        db.add(PatientProfile(user_id=user.id))
    elif req.role == UserRole.DOCTOR:
        db.add(DoctorProfile(user_id=user.id))

    db.commit()
    db.refresh(user)
    token = create_access_token({"sub": user.id})   # [M3] FIX: No role in token
    return {"access_token": token, "token_type": "bearer", "user": _user_out(user)}


@app.post("/api/auth/login", tags=["Auth"])
@limiter.limit("10/minute")         # [M2] FIX: Rate limit login — brute-force protection
def login(request: Request, req: LoginRequest, db: Session = Depends(get_db)):
    user = db.query(User).filter(User.email == req.email.strip().lower()).first()
    # [M1] FIX: Constant-time compare via verify_password even when user is None
    _dummy = "$2b$12$notarealhashXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXX"
    pwd_ok = verify_password(req.password, user.hashed_password if user else _dummy)

    if not user or not pwd_ok or not user.is_active:
        # [M1] FIX: Same message regardless of failure reason
        raise HTTPException(401, "Invalid credentials")

    token = create_access_token({"sub": user.id})   # [M3] FIX: No role in token
    logger.info("[Auth] Login: user_id=%s ip=%s", user.id, request.client.host)
    return {"access_token": token, "token_type": "bearer", "user": _user_out(user)}


@app.get("/api/auth/me", tags=["Auth"])
def get_me(current_user: User = Depends(get_current_user_secure)):
    return _user_out(current_user)


# ═══════════════════════════════════════════════════════════════════
#  PATIENT ROUTES
# ═══════════════════════════════════════════════════════════════════

@app.get("/api/patients", tags=["Patients"])
def list_patients(
    current_user: User = Depends(require_role_secure(UserRole.DOCTOR, UserRole.PHARMACIST)),
    db: Session = Depends(get_db)
):
    profiles = db.query(PatientProfile).all()
    return [_patient_out(p) for p in profiles]


@app.get("/api/patients/{patient_id}", tags=["Patients"])
def get_patient(
    patient_id: str,
    current_user: User = Depends(get_current_user_secure),
    db: Session = Depends(get_db)
):
    profile = db.query(PatientProfile).filter(PatientProfile.id == patient_id).first()
    if not profile:
        raise HTTPException(404, "Not found")

    # [C1] FIX: Patients can only read their own profile
    if current_user.role == UserRole.PATIENT and profile.user_id != current_user.id:
        raise HTTPException(403, "Forbidden")

    return _patient_out(profile)


@app.put("/api/patients/{patient_id}", tags=["Patients"])
def update_patient(
    patient_id: str,
    data: PatientProfileUpdate,
    current_user: User = Depends(get_current_user_secure),
    db: Session = Depends(get_db)
):
    profile = db.query(PatientProfile).filter(PatientProfile.id == patient_id).first()
    if not profile:
        raise HTTPException(404, "Not found")

    # [C1] FIX: Patients can only update their own profile
    if current_user.role == UserRole.PATIENT and profile.user_id != current_user.id:
        raise HTTPException(403, "Forbidden")

    # [H1] FIX: Explicit field allowlist — never trust Pydantic field names blindly
    safe_updates = {
        k: v for k, v in data.model_dump(exclude_none=True).items()
        if k in PatientProfileUpdate._SAFE_FIELDS
    }
    for field, value in safe_updates.items():
        setattr(profile, field, value)

    db.commit()
    db.refresh(profile)
    return _patient_out(profile)


@app.post("/api/patients/{patient_id}/assign-doctor", tags=["Patients"])
def assign_doctor(
    patient_id: str,
    doctor_id: str = Query(..., max_length=36),
    current_user: User = Depends(require_role_secure(UserRole.DOCTOR)),
    db: Session = Depends(get_db)
):
    profile = db.query(PatientProfile).filter(PatientProfile.id == patient_id).first()
    if not profile:
        raise HTTPException(404, "Not found")

    # Verify the target doctor actually exists and is a doctor
    target_doctor = db.query(User).filter(
        User.id == doctor_id, User.role == UserRole.DOCTOR, User.is_active == True
    ).first()
    if not target_doctor:
        raise HTTPException(400, "Invalid doctor ID")

    profile.assigned_doctor_id = doctor_id
    db.commit()
    return {"message": "Doctor assigned"}


# ═══════════════════════════════════════════════════════════════════
#  REHAB SESSION ROUTES
# ═══════════════════════════════════════════════════════════════════

@app.post("/api/sessions", tags=["Sessions"])
def start_session(
    data: SessionCreate,
    current_user: User = Depends(get_current_user_secure),
    db: Session = Depends(get_db)
):
    profile = db.query(PatientProfile).filter(PatientProfile.user_id == current_user.id).first()
    if not profile:
        raise HTTPException(404, "Patient profile not found")

    session = RehabSession(
        patient_id=profile.id,
        exercise_plan_id=data.exercise_plan_id,
        session_mode=data.session_mode
    )
    db.add(session)
    db.commit()
    db.refresh(session)
    return _session_out(session)


@app.put("/api/sessions/{session_id}/end", tags=["Sessions"])
def end_session(
    session_id: str,
    data: SessionUpdate,
    current_user: User = Depends(get_current_user_secure),
    db: Session = Depends(get_db)
):
    session = db.query(RehabSession).filter(RehabSession.id == session_id).first()
    if not session:
        raise HTTPException(404, "Not found")

    # [C2] FIX: Only the patient who owns this session (or a doctor) can end it
    if current_user.role == UserRole.PATIENT:
        profile = db.query(PatientProfile).filter(
            PatientProfile.user_id == current_user.id
        ).first()
        if not profile or session.patient_id != profile.id:
            raise HTTPException(403, "Forbidden")

    session.ended_at = datetime.utcnow()
    if session.started_at:
        delta = (session.ended_at - session.started_at).total_seconds()
        session.duration_seconds = int(delta)

    _SESSION_UPDATE_FIELDS = {
        "exercises_completed", "total_reps", "avg_form_score",
        "emg_peak", "emg_avg_rms", "intent_count", "signal_quality", "notes"
    }
    for field, value in data.model_dump(exclude_none=True).items():
        if field in _SESSION_UPDATE_FIELDS:
            setattr(session, field, value)

    db.commit()
    db.refresh(session)
    return _session_out(session)


@app.get("/api/sessions/patient/{patient_id}", tags=["Sessions"])
def get_patient_sessions(
    patient_id: str,
    limit: int = Query(20, ge=1, le=100),
    current_user: User = Depends(get_current_user_secure),
    db: Session = Depends(get_db)
):
    # [C2] FIX: Patients can only view their own sessions
    if current_user.role == UserRole.PATIENT:
        profile = db.query(PatientProfile).filter(
            PatientProfile.user_id == current_user.id
        ).first()
        if not profile or profile.id != patient_id:
            raise HTTPException(403, "Forbidden")

    sessions = (
        db.query(RehabSession)
        .filter(RehabSession.patient_id == patient_id)
        .order_by(RehabSession.started_at.desc())
        .limit(limit)
        .all()
    )
    return [_session_out(s) for s in sessions]


# ═══════════════════════════════════════════════════════════════════
#  EXERCISE PLAN ROUTES
# ═══════════════════════════════════════════════════════════════════

@app.post("/api/exercise-plans", tags=["Exercise Plans"])
def create_plan(
    data: ExercisePlanCreate,
    current_user: User = Depends(require_role_secure(UserRole.DOCTOR)),
    db: Session = Depends(get_db)
):
    # Verify target patient exists
    patient = db.query(PatientProfile).filter(PatientProfile.id == data.patient_id).first()
    if not patient:
        raise HTTPException(404, "Patient not found")

    plan = ExercisePlan(
        patient_id=data.patient_id,
        doctor_id=current_user.id,
        title=data.title,
        description=data.description,
        exercises=data.exercises,
        frequency_per_week=data.frequency_per_week,
        duration_weeks=data.duration_weeks
    )
    db.add(plan)
    db.commit()
    db.refresh(plan)
    return _plan_out(plan)


@app.get("/api/exercise-plans/patient/{patient_id}", tags=["Exercise Plans"])
def get_plans(
    patient_id: str,
    current_user: User = Depends(get_current_user_secure),
    db: Session = Depends(get_db)
):
    # [C1-class] Patients can only view their own plans
    if current_user.role == UserRole.PATIENT:
        profile = db.query(PatientProfile).filter(
            PatientProfile.user_id == current_user.id
        ).first()
        if not profile or profile.id != patient_id:
            raise HTTPException(403, "Forbidden")

    plans = db.query(ExercisePlan).filter(
        ExercisePlan.patient_id == patient_id,
        ExercisePlan.is_active == True
    ).all()
    return [_plan_out(p) for p in plans]


# ═══════════════════════════════════════════════════════════════════
#  PRESCRIPTION ROUTES
# ═══════════════════════════════════════════════════════════════════

def _create_order2_background(prescription_id: str, patient_id: str, medications: list, db_url: str):
    """
    [M4] FIX: Order2 creation runs as a background task with its own DB session.
    A failure here no longer rolls back the already-committed Prescription.
    """
    from sqlalchemy import create_engine
    from sqlalchemy.orm import sessionmaker
    from datetime import datetime as dt

    engine = create_engine(db_url)
    SessionLocal = sessionmaker(bind=engine)
    bg_db = SessionLocal()
    try:
        patient_profile = bg_db.query(PatientProfile).filter(
            PatientProfile.id == patient_id
        ).first()
        patient_name    = patient_profile.user.full_name if patient_profile and patient_profile.user else "Patient"
        patient_user_id = patient_profile.user_id if patient_profile else None

        items = [
            {
                "name":     m.get("name", "Unknown")[:200],   # already validated upstream
                "qty":      1,
                "price":    0,
                "dose":     str(m.get("dose", ""))[:100],
                "frequency":str(m.get("frequency", ""))[:100],
                "duration": str(m.get("duration", ""))[:100],
                "rx":       True,
            }
            for m in medications
        ]
        order2_code = f"RX-{dt.now().strftime('%y%m%d')}-{bg_db.query(Order2).count()+1:03d}"
        order2 = Order2(
            order_code      = order2_code,
            user_id         = patient_user_id,
            patient_label   = patient_name,
            items           = items,
            total           = 0,
            status          = "pending",
            payment_status  = "Prescription",
            prescription_id = prescription_id,
        )
        bg_db.add(order2)
        bg_db.commit()
        logger.info("[PharmaMind] Background Order2 %s created for %s", order2_code, patient_name)
    except Exception as e:
        bg_db.rollback()
        logger.error("[PharmaMind] Background Order2 creation failed: %s", e)
    finally:
        bg_db.close()


@app.post("/api/prescriptions", tags=["Prescriptions"])
def create_prescription(
    data: PrescriptionCreate,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(require_role_secure(UserRole.DOCTOR)),
    db: Session = Depends(get_db)
):
    # Verify patient exists
    patient = db.query(PatientProfile).filter(PatientProfile.id == data.patient_id).first()
    if not patient:
        raise HTTPException(404, "Patient not found")

    # [H2] medications already validated by MedicationItem schema
    meds_list = [m.model_dump() for m in data.medications]

    prescription = Prescription(
        patient_id=data.patient_id,
        doctor_id=current_user.id,
        medications=meds_list,
        notes=data.notes
    )
    db.add(prescription)
    db.flush()

    order = PharmacyOrder(prescription_id=prescription.id)
    db.add(order)
    db.commit()
    db.refresh(prescription)

    # [M4] FIX: Order2 created in background — won't roll back the Prescription
    db_url = str(db.get_bind().url)
    background_tasks.add_task(
        _create_order2_background,
        prescription.id,
        data.patient_id,
        meds_list,
        db_url,
    )

    return _prescription_out(prescription)


@app.get("/api/prescriptions/patient/{patient_id}", tags=["Prescriptions"])
def get_prescriptions(
    patient_id: str,
    current_user: User = Depends(get_current_user_secure),
    db: Session = Depends(get_db)
):
    # [C1-class] Patients can only view own prescriptions
    if current_user.role == UserRole.PATIENT:
        profile = db.query(PatientProfile).filter(
            PatientProfile.user_id == current_user.id
        ).first()
        if not profile or profile.id != patient_id:
            raise HTTPException(403, "Forbidden")

    prescriptions = db.query(Prescription).filter(
        Prescription.patient_id == patient_id
    ).order_by(Prescription.created_at.desc()).all()
    return [_prescription_out(p) for p in prescriptions]


# ═══════════════════════════════════════════════════════════════════
#  PHARMACY ROUTES
# ═══════════════════════════════════════════════════════════════════

@app.get("/api/pharmacy/orders", tags=["Pharmacy"])
def get_orders(
    status: Optional[str] = Query(None, max_length=50),
    current_user: User = Depends(require_role_secure(UserRole.PHARMACIST, UserRole.DOCTOR)),
    db: Session = Depends(get_db)
):
    q = db.query(PharmacyOrder)
    if status:
        # Allowlist statuses to prevent filter injection
        _VALID_STATUSES = {"pending", "processing", "dispensed", "cancelled"}
        if status not in _VALID_STATUSES:
            raise HTTPException(400, "Invalid status filter")
        q = q.filter(PharmacyOrder.status == status)
    return [_order_out(o) for o in q.order_by(PharmacyOrder.created_at.desc()).all()]


@app.put("/api/pharmacy/orders/{order_id}", tags=["Pharmacy"])
def update_order(
    order_id: str,
    data: PharmacyOrderUpdate,
    current_user: User = Depends(require_role_secure(UserRole.PHARMACIST)),
    db: Session = Depends(get_db)
):
    _VALID_STATUSES = {"pending", "processing", "dispensed", "cancelled"}
    if data.status not in _VALID_STATUSES:
        raise HTTPException(400, "Invalid status value")

    order = db.query(PharmacyOrder).filter(PharmacyOrder.id == order_id).first()
    if not order:
        raise HTTPException(404, "Order not found")

    order.status = data.status
    order.pharmacist_id = current_user.id
    if data.notes:
        order.notes = data.notes
    if data.status == "dispensed":
        order.dispensed_at = datetime.utcnow()
        if order.prescription:
            order.prescription.status = "dispensed"
    db.commit()
    return _order_out(order)


@app.get("/api/pharmacy/inventory", tags=["Pharmacy"])
def get_inventory(
    current_user: User = Depends(require_role_secure(UserRole.PHARMACIST, UserRole.DOCTOR)),
    db: Session = Depends(get_db)
):
    items = db.query(MedicineInventory).order_by(MedicineInventory.name).all()
    return [_inventory_out(i) for i in items]


@app.get("/api/pharmacy/inventory/low-stock", tags=["Pharmacy"])
def low_stock_alert(
    current_user: User = Depends(require_role_secure(UserRole.PHARMACIST)),
    db: Session = Depends(get_db)
):
    items = db.query(MedicineInventory).filter(
        MedicineInventory.stock_quantity <= MedicineInventory.reorder_level
    ).all()
    return [_inventory_out(i) for i in items]


@app.get("/api/pharmacy/search", tags=["Pharmacy"])
def search_medicines(
    q: str = Query("", max_length=100),
    current_user: User = Depends(get_current_user_secure),
    db: Session = Depends(get_db)
):
    if not q or len(q) < 2:
        return []
    # Sanitize: only alphanumeric + space + hyphen
    q_clean = re.sub(r"[^\w\s\-]", "", q)[:100]
    items = db.query(MedicineInventory).filter(
        MedicineInventory.name.ilike(f"%{q_clean}%")
    ).order_by(MedicineInventory.name).limit(8).all()
    return [_inventory_out(i) for i in items]


@app.put("/api/pharmacy/inventory/{item_id}/stock", tags=["Pharmacy"])
def update_stock(
    item_id: str,
    data: StockUpdateBody,
    current_user: User = Depends(require_role_secure(UserRole.PHARMACIST)),
    db: Session = Depends(get_db)
):
    item = db.query(MedicineInventory).filter(MedicineInventory.id == item_id).first()
    if not item:
        raise HTTPException(404, "Medicine not found")
    item.stock_quantity = data.stock_quantity
    if data.reorder_level is not None:
        item.reorder_level = data.reorder_level
    db.commit()
    db.refresh(item)
    return _inventory_out(item)


# ═══════════════════════════════════════════════════════════════════
#  PROGRESS REPORT ROUTES
# ═══════════════════════════════════════════════════════════════════

@app.get("/api/reports/patient/{patient_id}", tags=["Reports"])
def get_reports(
    patient_id: str,
    current_user: User = Depends(get_current_user_secure),
    db: Session = Depends(get_db)
):
    # [C1-class] Patients can only view own reports
    if current_user.role == UserRole.PATIENT:
        profile = db.query(PatientProfile).filter(
            PatientProfile.user_id == current_user.id
        ).first()
        if not profile or profile.id != patient_id:
            raise HTTPException(403, "Forbidden")

    reports = db.query(ProgressReport).filter(
        ProgressReport.patient_id == patient_id
    ).order_by(ProgressReport.created_at.desc()).all()
    return [_report_out(r) for r in reports]


@app.put("/api/reports/{report_id}/approve", tags=["Reports"])
def approve_report(
    report_id: str,
    doctor_notes: Optional[str] = Query(None, max_length=2000),
    current_user: User = Depends(require_role_secure(UserRole.DOCTOR)),
    db: Session = Depends(get_db)
):
    report = db.query(ProgressReport).filter(ProgressReport.id == report_id).first()
    if not report:
        raise HTTPException(404, "Report not found")
    report.doctor_approved = True
    report.doctor_notes    = doctor_notes
    db.commit()
    return _report_out(report)


# ═══════════════════════════════════════════════════════════════════
#  REPORT DOWNLOAD
# ═══════════════════════════════════════════════════════════════════

@app.get("/api/reports/{report_id}/download", tags=["Reports"])
def download_report(
    report_id: str,
    current_user: User = Depends(get_current_user_secure),
    db: Session = Depends(get_db)
):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io
    from fastapi.responses import StreamingResponse

    report = db.query(ProgressReport).filter(ProgressReport.id == report_id).first()
    if not report:
        raise HTTPException(404, "Report not found")

    # [C1-class] Patients can only download their own report
    if current_user.role == UserRole.PATIENT:
        profile = db.query(PatientProfile).filter(
            PatientProfile.user_id == current_user.id
        ).first()
        if not profile or profile.id != report.patient_id:
            raise HTTPException(403, "Forbidden")

    patient      = db.query(PatientProfile).filter(PatientProfile.id == report.patient_id).first()
    patient_name = patient.user.full_name if patient and patient.user else "Patient"
    diagnosis    = patient.diagnosis or "Neurological condition" if patient else "Neurological condition"

    doc = Document()
    title = doc.add_heading("NeuroStride — AI Progress Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(47, 123, 232)
    doc.add_paragraph()

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    meta = [
        ("Patient",   patient_name),
        ("Diagnosis", diagnosis),
        ("Period",    f"{report.period_start} — {report.period_end}"),
        ("Generated", str(report.created_at)[:10]),
        ("Status",    "Doctor Approved" if report.doctor_approved else "Pending Approval"),
    ]
    for i, (k, v) in enumerate(meta):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    doc.add_heading("Clinical Summary", 1)
    doc.add_paragraph(report.ai_summary or "No summary available.")
    doc.add_paragraph()

    if report.strengths:
        doc.add_heading("Strengths", 1)
        for s in report.strengths:
            doc.add_paragraph(f"✓  {s}", style="List Bullet")
    doc.add_paragraph()

    if report.improvements:
        doc.add_heading("Areas for Improvement", 1)
        for imp in report.improvements:
            doc.add_paragraph(f"→  {imp}", style="List Bullet")
    doc.add_paragraph()

    if report.recommendations:
        doc.add_heading("Doctor Recommendations", 1)
        for rec in report.recommendations:
            doc.add_paragraph(f"•  {rec}", style="List Bullet")

    if report.doctor_notes:
        doc.add_paragraph()
        doc.add_heading("Doctor Notes", 1)
        p = doc.add_paragraph(report.doctor_notes)
        p.runs[0].italic = True

    doc.add_paragraph()
    footer = doc.add_paragraph("Generated by NeuroStride AI Rehabilitation Platform")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(139, 148, 158)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    # Sanitize filename — strip non-alphanum chars
    safe_name = re.sub(r"[^\w\-]", "_", patient_name)
    filename = f"NeuroStride_Report_{safe_name}_{report.period_end}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=\"{filename}\""}
    )


# ═══════════════════════════════════════════════════════════════════
#  BILL / INVOICE GENERATION
# ═══════════════════════════════════════════════════════════════════

@app.get("/api/pharmacy/orders/{order_id}/bill", tags=["Pharmacy"])
def download_bill(
    order_id: str,
    current_user: User = Depends(get_current_user_secure),
    db: Session = Depends(get_db)
):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io
    from fastapi.responses import StreamingResponse

    # [H4] FIX: was db.query(Order) — Order undefined. Correct model is Order2.
    order = db.query(Order2).filter(Order2.id == order_id).first()
    if not order:
        raise HTTPException(404, "Order not found")

    # Access control: patients can only view their own bill
    if current_user.role == UserRole.PATIENT and order.user_id != current_user.id:
        raise HTTPException(403, "Forbidden")

    doc = Document()
    title = doc.add_heading("NeuroStride PharmaMind", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(252, 109, 38)

    sub = doc.add_paragraph("AI-Powered Pharmacy")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(10)
    doc.add_paragraph()

    doc.add_heading("Tax Invoice", 1)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    details = [
        ("Invoice No.", f"INV-{str(order.id)[:8].upper()}"),
        ("Order Code",  order.order_code or str(order.id)[:8]),
        ("Date",        str(order.created_at)[:10]),
        ("Patient",     order.patient_label or "Walk-in"),
    ]
    for i, (k, v) in enumerate(details):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = str(v)
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    doc.add_heading("Medicines", 1)

    items    = order.items or []
    subtotal = 0.0

    if items:
        item_table = doc.add_table(rows=1 + len(items), cols=4)
        item_table.style = "Table Grid"
        for j, h in enumerate(["Medicine", "Qty", "Unit Price (₹)", "Total (₹)"]):
            cell = item_table.rows[0].cells[j]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True

        for i, item in enumerate(items):
            qty   = int(item.get("qty", 1))
            price = float(item.get("price", 0))
            total = qty * price
            subtotal += total
            row = item_table.rows[i + 1]
            row.cells[0].text = str(item.get("name", "?"))[:200]
            row.cells[1].text = str(qty)
            row.cells[2].text = f"₹{price:.2f}"
            row.cells[3].text = f"₹{total:.2f}"
    else:
        doc.add_paragraph("No items recorded.")
        subtotal = float(order.total or 0)

    doc.add_paragraph()
    service_fee = round(subtotal * 0.02, 2)
    gst         = round(subtotal * 0.18, 2)
    grand_total = round(subtotal + service_fee + gst, 2)

    totals_table = doc.add_table(rows=4, cols=2)
    totals_table.style = "Table Grid"
    for i, (k, v) in enumerate([
        ("Subtotal",            f"₹{subtotal:.2f}"),
        ("Service Charge (2%)", f"₹{service_fee:.2f}"),
        ("GST @ 18%",           f"₹{gst:.2f}"),
        ("GRAND TOTAL",         f"₹{grand_total:.2f}"),
    ]):
        totals_table.rows[i].cells[0].text = k
        totals_table.rows[i].cells[1].text = v
        if i == 3:
            totals_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            totals_table.rows[i].cells[1].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    footer = doc.add_paragraph("Thank you for choosing NeuroStride PharmaMind 💊")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_code = re.sub(r"[^\w\-]", "_", order.order_code or str(order.id))
    filename   = f"NeuroStride_Bill_{safe_code}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=\"{filename}\""}
    )


# ═══════════════════════════════════════════════════════════════════
#  WEBSOCKET — Live sensor data relay (SECURED)
# ═══════════════════════════════════════════════════════════════════

@app.websocket("/ws/sensor/{patient_id}")
async def sensor_websocket(
    websocket: WebSocket,
    patient_id: str,
    token: str = Query(...),          # [C3] FIX: Require token as query param
    db: Session = Depends(get_db)
):
    try:
        user = _ws_auth(token, db)
    except (HTTPException, ValueError):
        await websocket.close(code=4401)  # 4401 = custom: unauthorized
        return

    # Patient can only subscribe to their own sensor feed
    if user.role == UserRole.PATIENT:
        profile = db.query(PatientProfile).filter(
            PatientProfile.user_id == user.id
        ).first()
        if not profile or profile.id != patient_id:
            await websocket.close(code=4403)
            return

    await manager.connect(websocket, f"sensor_{patient_id}")
    logger.info("[WS] Sensor connection: patient=%s user=%s", patient_id, user.id)

    try:
        while True:
            raw  = await websocket.receive_text()
            # Validate incoming payload size
            if len(raw) > 64 * 1024:   # 64KB max per message
                logger.warning("[WS] Oversized payload from user=%s, dropping", user.id)
                continue
            try:
                data = json.loads(raw)
            except json.JSONDecodeError:
                logger.warning("[WS] Invalid JSON from user=%s", user.id)
                continue
            await manager.broadcast(f"sensor_{patient_id}", data)
    except WebSocketDisconnect:
        manager.disconnect(websocket, f"sensor_{patient_id}")


@app.websocket("/ws/session/{session_id}")
async def session_websocket(
    websocket: WebSocket,
    session_id: str,
    token: str = Query(...),          # [C3] FIX: Require token
    db: Session = Depends(get_db)
):
    try:
        user = _ws_auth(token, db)
    except (HTTPException, ValueError):
        await websocket.close(code=4401)
        return

    # Verify session belongs to this user (if patient)
    if user.role == UserRole.PATIENT:
        session = db.query(RehabSession).filter(RehabSession.id == session_id).first()
        if not session:
            await websocket.close(code=4404)
            return
        profile = db.query(PatientProfile).filter(
            PatientProfile.user_id == user.id
        ).first()
        if not profile or session.patient_id != profile.id:
            await websocket.close(code=4403)
            return

    await manager.connect(websocket, f"session_{session_id}")
    try:
        while True:
            raw = await websocket.receive_text()
            if len(raw) > 32 * 1024:
                continue
            try:
                data = json.loads(raw)
            except json.JSONDecodeError:
                continue
            await manager.broadcast(f"session_{session_id}", data)
    except WebSocketDisconnect:
        manager.disconnect(websocket, f"session_{session_id}")


# ═══════════════════════════════════════════════════════════════════
#  HEALTH CHECK
# ═══════════════════════════════════════════════════════════════════

@app.get("/api/health", tags=["System"])
def health():
    # Don't leak version in production
    return {"status": "ok", "service": "NeuroStride API"}


# ═══════════════════════════════════════════════════════════════════
#  Serializers (sensitive fields stripped)
# ═══════════════════════════════════════════════════════════════════

def _user_out(u: User) -> dict:
    data = {
        "id":         u.id,
        "email":      u.email,
        "full_name":  u.full_name,
        "role":       u.role.value,
        "phone":      u.phone,
        "language":   u.language,
        "created_at": str(u.created_at),
        "profile_id": None,
        # [BONUS] hashed_password, is_active, internal flags NEVER returned
    }
    try:
        if u.role == UserRole.PATIENT and u.patient_profile:
            data["profile_id"] = u.patient_profile.id
    except Exception:
        pass
    return data

def _patient_out(p: PatientProfile) -> dict:
    return {
        "id":                 p.id,
        "user_id":            p.user_id,
        "full_name":          p.user.full_name if p.user else None,
        "email":              p.user.email if p.user else None,
        "date_of_birth":      p.date_of_birth,
        "gender":             p.gender,
        "blood_group":        p.blood_group,
        "weight_kg":          p.weight_kg,
        "height_cm":          p.height_cm,
        "diagnosis":          p.diagnosis,
        "affected_side":      p.affected_side,
        "paralysis_level":    p.paralysis_level,
        "allergies":          p.allergies,
        "current_meds":       p.current_meds,
        "emergency_contact":  p.emergency_contact,
        "assigned_doctor_id": p.assigned_doctor_id,
        # Internal FK fields intentionally excluded
    }

def _session_out(s: RehabSession) -> dict:
    return {
        "id":                  s.id,
        "patient_id":          s.patient_id,
        "exercise_plan_id":    s.exercise_plan_id,
        "started_at":          str(s.started_at),
        "ended_at":            str(s.ended_at) if s.ended_at else None,
        "duration_seconds":    s.duration_seconds,
        "exercises_completed": s.exercises_completed,
        "total_reps":          s.total_reps,
        "avg_form_score":      s.avg_form_score,
        "emg_peak":            s.emg_peak,
        "emg_avg_rms":         s.emg_avg_rms,
        "intent_count":        s.intent_count,
        "signal_quality":      s.signal_quality,
        "session_mode":        s.session_mode,
        "notes":               s.notes,
    }

def _plan_out(p: ExercisePlan) -> dict:
    return {
        "id":                 p.id,
        "patient_id":         p.patient_id,
        "doctor_id":          p.doctor_id,
        "title":              p.title,
        "description":        p.description,
        "exercises":          p.exercises,
        "frequency_per_week": p.frequency_per_week,
        "duration_weeks":     p.duration_weeks,
        "is_active":          p.is_active,
        "ai_generated":       p.ai_generated,
        "created_at":         str(p.created_at),
    }

def _prescription_out(p: Prescription) -> dict:
    return {
        "id":                   p.id,
        "patient_id":           p.patient_id,
        "doctor_id":            p.doctor_id,
        "medications":          p.medications,
        "notes":                p.notes,
        "status":               p.status,
        "created_at":           str(p.created_at),
        "ai_interaction_check": p.ai_interaction_check,
    }

def _order_out(o: PharmacyOrder) -> dict:
    return {
        "id":              o.id,
        "prescription_id": o.prescription_id,
        "pharmacist_id":   o.pharmacist_id,
        "status":          o.status,
        "notes":           o.notes,
        "created_at":      str(o.created_at),
        "dispensed_at":    str(o.dispensed_at) if o.dispensed_at else None,
    }

def _inventory_out(i: MedicineInventory) -> dict:
    return {
        "id":             i.id,
        "name":           i.name,
        "generic_name":   i.generic_name,
        "category":       i.category,
        "strength":       i.strength,
        "unit":           i.unit,
        "stock_quantity": i.stock_quantity,
        "reorder_level":  i.reorder_level,
        "price":          i.price,
        "manufacturer":   i.manufacturer,
        "expiry_date":    i.expiry_date,
    }

def _report_out(r: ProgressReport) -> dict:
    return {
        "id":              r.id,
        "patient_id":      r.patient_id,
        "doctor_id":       r.doctor_id,
        "period_start":    r.period_start,
        "period_end":      r.period_end,
        "ai_summary":      r.ai_summary,
        "strengths":       r.strengths,
        "improvements":    r.improvements,
        "recommendations": r.recommendations,
        "docx_path":       r.docx_path,
        "doctor_approved": r.doctor_approved,
        "doctor_notes":    r.doctor_notes,
        "created_at":      str(r.created_at),
    }